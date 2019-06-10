# -*- coding: utf-8 -*-
from datetime import datetime
import os

from docx import Document
from docx.shared import RGBColor
from docx.shared import Inches
from docx.shared import Pt

class SimpleDocxService:

    def __init__(self, filename=None):
        self.document = Document(filename)
        self.paragraphs = self.document.paragraphs
        self.filename = filename

    @property
    def temp_fold(self):
        if not os.path.exists(self.filename):
            os.mkdir(self.filename)
        return self.filename

    def set_normal_font(self, name, size):
        """フォントの設定"""
        font = self.document.styles['Normal'].font
        font.name = name
        font.size = Pt(size)

    def add_head(self, text, lv):
        """見出しの設定"""
        self.document.add_heading(text, level=lv)

    def fill_information(self, user_name, questions, similar, create_time=datetime.now().strftime("%Y-%m-%d")):
        self.paragraphs[2].text = self.paragraphs[2].text.replace('{user_name}', user_name).replace('{create_time}', create_time)
        self.paragraphs[3].text = self.paragraphs[3].text.replace('{count}', questions).replace('{similar}', similar)

    def open_text(self):
        """テキスト追加開始"""
        class Paragraph:

            def __init__(self, paragraph):
                self.paragraph = paragraph
                self.text = None

            def __enter__(self):
                """テキスト追加開始"""
                return self

            def add(self, text):
                """テキスト追加"""
                self.text = self.paragraph.add_run(text)
                return self

            def italic(self):
                """斜体にする"""
                self.text.italic = True
                return self

            def bold(self):
                """太字にする"""
                self.text.bold = True
                return self

            def color(self, r, g, b):
                """色を付ける"""
                self.text.font.color.rgb = RGBColor(r, g, b)
                return self

            def close(self):
                """テキスト追加終了"""
                del self.paragraph
                del self.text

            def __exit__(self, *args):
                """テキスト追加終了"""
                self.close()

        return Paragraph(self.document.add_paragraph())

    def add_picture(self, filename, inch):
        """図を挿入する"""
        self.document.add_picture(filename, width=Inches(inch))

    def save(self, name):
        """docxファイルとして出力"""
        self.document.save(name)